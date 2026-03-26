"""
Form Filler — reads .docx bank forms, sends to Claude with the data dictionary,
gets back field mappings, and fills the form using python-docx.
"""

import anthropic
import base64
import json
import logging
import re
import copy
from io import BytesIO
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT

logger = logging.getLogger(__name__)

client = anthropic.Anthropic()

FILL_PROMPT_RU = """Ты заполняешь банковскую форму для компании.

СПРАВОЧНИК ДАННЫХ (источник данных):
{dictionary}

ТЕКСТ БАНКОВСКОЙ ФОРМЫ (поля для заполнения):
{form_text}

Инструкции:
1. Определи КАЖДОЕ заполняемое поле в форме (пустые ячейки, подчёркивания, скобки, поля с подсказками)
2. Сопоставь каждое поле с правильным значением из справочника
3. Будь гибким: названия полей в формах разных банков отличаются —
   «Наименование организации», «Название компании», «Полное наименование» — это одно и то же
4. Для полей, которых нет в справочнике, используй null
5. Обрати внимание: если поле спрашивает про бенефициарного владельца (UBO), это акционер с долей >25%
6. Язык значений: русский

Верни JSON-массив:
[
  {{"field_label": "точный текст метки поля из формы", "value": "значение для вставки", "table_index": номер_таблицы_или_null, "row_index": номер_строки_или_null, "col_index": номер_столбца_или_null}},
  ...
]
Верни ТОЛЬКО валидный JSON без markdown."""

FILL_PROMPT_EN = """You are filling out a bank form for a company.

DATA DICTIONARY (source of truth):
{dictionary}

BANK FORM TEXT (fields to fill):
{form_text}

Instructions:
1. Identify EVERY fillable field in the form (empty cells, underscores, brackets, placeholder text)
2. Match each field to the correct value from the dictionary
3. Be flexible: field labels vary across banks —
   "Company Name", "Legal Name", "Name of Organisation" all map to the company name
4. For fields not in the dictionary, use null
5. Note: if a field asks about beneficial owner (UBO), that's the shareholder with >25% ownership
6. Language for values: English

Return JSON array:
[
  {{"field_label": "exact text label from form", "value": "value to insert", "table_index": table_number_or_null, "row_index": row_number_or_null, "col_index": column_number_or_null}},
  ...
]
Return ONLY valid JSON with no markdown."""


def extract_form_text(doc: Document) -> str:
    """Extract all text from a .docx including paragraphs and tables with position info."""
    parts = []

    # Paragraphs
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            parts.append(f"[Paragraph {i}] {text}")

    # Tables with row/col indices
    for t_idx, table in enumerate(doc.tables):
        parts.append(f"\n[Table {t_idx}]")
        for r_idx, row in enumerate(table.rows):
            cells = []
            for c_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip() or "(empty)"
                cells.append(f"[col{c_idx}] {cell_text}")
            parts.append(f"  Row {r_idx}: {' | '.join(cells)}")

    return "\n".join(parts)


def get_field_mappings(form_text: str, dictionary: dict, language: str = "ru") -> list[dict]:
    """Send form text + dictionary to Claude, get back field mappings."""
    template = FILL_PROMPT_RU if language == "ru" else FILL_PROMPT_EN
    prompt = template.format(
        dictionary=json.dumps(dictionary, ensure_ascii=False, indent=2),
        form_text=form_text,
    )

    response = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=8192,
        messages=[{"role": "user", "content": prompt}],
    )

    raw = response.content[0].text.strip()
    if raw.startswith("```"):
        raw = raw.split("\n", 1)[1]
    if raw.endswith("```"):
        raw = raw.rsplit("```", 1)[0]
    raw = raw.strip()

    try:
        mappings = json.loads(raw)
        return [m for m in mappings if m.get("value") is not None]
    except json.JSONDecodeError as e:
        logger.error(f"Failed to parse mappings JSON: {e}\nRaw: {raw[:500]}")
        return []


def apply_mappings(doc: Document, mappings: list[dict]) -> Document:
    """
    Apply field mappings to the document.
    Strategy:
    1. For table-based mappings (with table/row/col indices), fill directly
    2. For text-based mappings, find the label and fill adjacent blanks
    """
    for mapping in mappings:
        value = str(mapping.get("value", ""))
        if not value or value == "null":
            continue

        label = mapping.get("field_label", "")
        t_idx = mapping.get("table_index")
        r_idx = mapping.get("row_index")
        c_idx = mapping.get("col_index")

        # Strategy 1: Direct table cell fill
        if t_idx is not None and r_idx is not None and c_idx is not None:
            try:
                table = doc.tables[t_idx]
                row = table.rows[r_idx]
                cell = row.cells[c_idx]
                cell_text = cell.text.strip()
                # Only fill if cell is empty or has placeholder
                if not cell_text or cell_text == "(empty)" or _is_placeholder(cell_text):
                    # Preserve formatting: write into first paragraph's first run
                    if cell.paragraphs and cell.paragraphs[0].runs:
                        cell.paragraphs[0].runs[0].text = value
                        # Clear other runs
                        for run in cell.paragraphs[0].runs[1:]:
                            run.text = ""
                    else:
                        cell.text = value
                    logger.debug(f"Filled table {t_idx} row {r_idx} col {c_idx}: {value}")
                    continue
            except (IndexError, AttributeError) as e:
                logger.warning(f"Table index error for {label}: {e}")

        # Strategy 2: Search tables for the label, fill adjacent cell
        filled = False
        for table in doc.tables:
            for row in table.rows:
                for i, cell in enumerate(row.cells):
                    if label and _text_matches(cell.text, label):
                        # Try filling next cell in the same row
                        if i + 1 < len(row.cells):
                            target = row.cells[i + 1]
                            target_text = target.text.strip()
                            if not target_text or _is_placeholder(target_text):
                                if target.paragraphs and target.paragraphs[0].runs:
                                    target.paragraphs[0].runs[0].text = value
                                    for run in target.paragraphs[0].runs[1:]:
                                        run.text = ""
                                else:
                                    target.text = value
                                filled = True
                                logger.debug(f"Filled by label match '{label}': {value}")
                                break
                if filled:
                    break
            if filled:
                break

        # Strategy 3: Search paragraphs for placeholder patterns
        if not filled and label:
            for para in doc.paragraphs:
                if label in para.text:
                    for run in para.runs:
                        if _is_placeholder(run.text):
                            run.text = value
                            filled = True
                            break
                    if filled:
                        break

    return doc


def _is_placeholder(text: str) -> bool:
    """Check if text is a placeholder (underscores, blanks, brackets)."""
    cleaned = text.strip()
    if not cleaned:
        return True
    if re.match(r'^[_\-\.]{2,}$', cleaned):
        return True
    if re.match(r'^\[[\s_]*\]$', cleaned):
        return True
    if cleaned in ("(empty)", "___", "______", "________"):
        return True
    return False


def _text_matches(cell_text: str, label: str) -> bool:
    """Fuzzy match a cell's text against a field label."""
    cell_clean = cell_text.strip().lower().rstrip(":").strip()
    label_clean = label.strip().lower().rstrip(":").strip()
    if not cell_clean or not label_clean:
        return False
    # Exact match
    if cell_clean == label_clean:
        return True
    # One contains the other
    if label_clean in cell_clean or cell_clean in label_clean:
        return True
    return False


async def fill_form(
    form_bytes: bytes,
    dictionary: dict,
    language: str = "ru",
) -> bytes:
    """
    Main entry point: takes .docx bytes + dictionary,
    returns filled .docx bytes.
    """
    doc = Document(BytesIO(form_bytes))
    form_text = extract_form_text(doc)

    logger.info(f"Extracted form text ({len(form_text)} chars), getting mappings...")
    mappings = get_field_mappings(form_text, dictionary, language)
    logger.info(f"Got {len(mappings)} field mappings, applying...")

    doc = apply_mappings(doc, mappings)

    # Save to bytes
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output.read()
